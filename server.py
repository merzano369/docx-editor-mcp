from mcp.server.fastmcp import FastMCP
from docx import Document
from docx.shared import Pt, Mm, RGBColor, Inches, Twips, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsmap
import os
import json
from datetime import datetime
from typing import Dict, Any, Optional, List, Union

# Initialize FastMCP server
mcp = FastMCP("docx-editor")
def _find_paragraph_index(doc: Document, target_text: str = None, target_paragraph_index: int = None) -> int:
    """
    Helper function to find a paragraph index either by text or by direct index.
    Returns the index, or -1 if not found.
    """
    if target_paragraph_index is not None:
        if 0 <= target_paragraph_index < len(doc.paragraphs):
            return target_paragraph_index
        return -1
        
    if target_text:
        for i, p in enumerate(doc.paragraphs):
            if target_text in p.text:
                return i
    return -1

@mcp.tool()
def insert_header_near_text(filename: str = None, target_text: str = None, header_title: str = None, position: str = 'after', header_style: str = 'Heading 1', target_paragraph_index: int = None) -> str:
    """
    Inserts a header relative to a specific paragraph or text match.
    """
    global current_doc
    doc = current_doc
    if filename and os.path.exists(filename):
        doc = Document(filename)
    if not doc:
        return "No active document. Call create_document first."
        
    if not header_title:
        return "header_title is required."

    p_idx = _find_paragraph_index(doc, target_text, target_paragraph_index)
    if p_idx == -1:
        return "Target paragraph not found."
        
    # Python-docx paragraph.insert_paragraph_before() inserts BEFORE the current paragraph.
    # To insert AFTER, we find the NEXT paragraph and insert before it. 
    # If it's the last paragraph, we just add_paragraph to the end of the document.
    
    target_p = doc.paragraphs[p_idx]
    
    if position == 'before':
        new_p = target_p.insert_paragraph_before(header_title, style=header_style)
    elif position == 'after':
        if p_idx + 1 < len(doc.paragraphs):
            next_p = doc.paragraphs[p_idx + 1]
            new_p = next_p.insert_paragraph_before(header_title, style=header_style)
        else:
            new_p = doc.add_paragraph(header_title, style=header_style)
    else:
        return "Invalid position. Use 'before' or 'after'."
        
    if filename:
        doc.save(filename)
    return f"Inserted header '{header_title}' {position} paragraph {p_idx}."

@mcp.tool()
def insert_line_or_paragraph_near_text(filename: str = None, target_text: str = None, line_text: str = None, position: str = 'after', line_style: str = None, target_paragraph_index: int = None) -> str:
    """
    Inserts a paragraph relative to a specific paragraph or text match.
    """
    global current_doc
    doc = current_doc
    if filename and os.path.exists(filename):
        doc = Document(filename)
    if not doc:
        return "No active document. Call create_document first."
        
    if not line_text:
        return "line_text is required."

    p_idx = _find_paragraph_index(doc, target_text, target_paragraph_index)
    if p_idx == -1:
        return "Target paragraph not found."
        
    target_p = doc.paragraphs[p_idx]
    
    if position == 'before':
        new_p = target_p.insert_paragraph_before(line_text, style=line_style)
    elif position == 'after':
        if p_idx + 1 < len(doc.paragraphs):
            next_p = doc.paragraphs[p_idx + 1]
            new_p = next_p.insert_paragraph_before(line_text, style=line_style)
        else:
            new_p = doc.add_paragraph(line_text, style=line_style)
    else:
        return "Invalid position. Use 'before' or 'after'."
        
    if filename:
        doc.save(filename)
    return f"Inserted paragraph {position} paragraph {p_idx}."

@mcp.tool()
def insert_numbered_list_near_text(filename: str = None, target_text: str = None, list_items: list = None, position: str = 'after', target_paragraph_index: int = None, bullet_type: str = 'bullet') -> str:
    """
    Inserts a list (bullet or numbered) relative to a specific paragraph or text match.
    """
    global current_doc
    doc = current_doc
    if filename and os.path.exists(filename):
        doc = Document(filename)
    if not doc:
        return "No active document. Call create_document first."
        
    if not list_items:
        return "list_items is required."

    p_idx = _find_paragraph_index(doc, target_text, target_paragraph_index)
    if p_idx == -1:
        return "Target paragraph not found."
        
    target_p = doc.paragraphs[p_idx]
    
    style_name = 'List Bullet' if bullet_type == 'bullet' else 'List Number'
    
    # We want to insert the items in order. 
    # If inserting 'before', we insert item 1 before target, then item 2 before target, etc.
    # Actually, inserting them sequentially before target means they would appear in reverse.
    # So we insert them exactly where needed.
    
    if position == 'before':
        # Insert all items before target_p, maintaining order
        for item in list_items:
            target_p.insert_paragraph_before(item, style=style_name)
    elif position == 'after':
        # To insert after and maintain order, we find the element AFTER target_p.
        if p_idx + 1 < len(doc.paragraphs):
            next_p = doc.paragraphs[p_idx + 1]
            for item in list_items:
                next_p.insert_paragraph_before(item, style=style_name)
        else:
            # If target_p is the last element, just append
            for item in list_items:
                doc.add_paragraph(item, style=style_name)
    else:
        return "Invalid position. Use 'before' or 'after'."
        
    if filename:
        doc.save(filename)
    return f"Inserted list of {len(list_items)} items {position} paragraph {p_idx}."

@mcp.tool()
def format_text(filename: str = None, paragraph_index: int = None, start_pos: int = 0, end_pos: int = None, bold: bool = None, italic: bool = None, underline: bool = None, color: str = None, font_size: int = None, font_name: str = None) -> str:
    """
    Formats a specific character range within a paragraph.
    Due to Word's internal representation, this clears the paragraph and rebuilds it with the requested formatting applied to the specific range.
    """
    global current_doc
    doc = current_doc
    if filename and os.path.exists(filename):
        doc = Document(filename)
    if not doc:
        return "No active document. Call create_document first."
        
    if paragraph_index is None or paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
        return "Invalid paragraph index."
        
    p = doc.paragraphs[paragraph_index]
    full_text = p.text
    
    if end_pos is None or end_pos > len(full_text):
        end_pos = len(full_text)
        
    if start_pos < 0 or start_pos >= end_pos:
        return "Invalid start_pos or end_pos."
        
    # We will clear the paragraph and rebuild it in 3 parts: before, target, after.
    # To preserve paragraph-level styles, we only delete the runs inside it.
    
    p.clear()
    
    part1_text = full_text[:start_pos]
    part2_text = full_text[start_pos:end_pos]
    part3_text = full_text[end_pos:]
    
    if part1_text:
        p.add_run(part1_text)
        
    if part2_text:
        run = p.add_run(part2_text)
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic
        if underline is not None:
            run.underline = underline
        if font_size is not None:
            run.font.size = Pt(font_size)
        if font_name is not None:
            run.font.name = font_name
        if color is not None:
            try:
                c = color.replace('#', '')
                if len(c) == 6:
                    run.font.color.rgb = RGBColor(int(c[0:2], 16), int(c[2:4], 16), int(c[4:6], 16))
            except:
                pass
                
    if part3_text:
        p.add_run(part3_text)
        
    if filename:
        doc.save(filename)
    return f"Formatted text in paragraph {paragraph_index} from position {start_pos} to {end_pos}."

@mcp.tool()
def search_and_replace(filename: str = None, find_text: str = None, replace_text: str = None) -> str:
    """
    Replaces all occurrences of find_text with replace_text across all document paragraphs.
    """
    global current_doc
    doc = current_doc
    if filename and os.path.exists(filename):
        doc = Document(filename)
    if not doc:
        return "No active document. Call create_document first."
        
    if not find_text:
        return "find_text is required."
        
    if replace_text is None:
        replace_text = ""
        
    replace_count = 0
    
    # Process all normal paragraphs
    for p in doc.paragraphs:
        if find_text in p.text:
            text = p.text
            new_text = text.replace(find_text, replace_text)
            p.clear()
            p.add_run(new_text)
            replace_count += 1
            
    # Process all table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if find_text in p.text:
                        text = p.text
                        new_text = text.replace(find_text, replace_text)
                        p.clear()
                        p.add_run(new_text)
                        replace_count += 1
                        
    if filename:
        doc.save(filename)
    return f"Replaced '{find_text}' with '{replace_text}' in {replace_count} elements."

@mcp.tool()
def delete_paragraph(filename: str = None, paragraph_index: int = None) -> str:
    """
    Deletes a paragraph by its internal index.
    """
    global current_doc
    doc = current_doc
    if filename and os.path.exists(filename):
        doc = Document(filename)
    if not doc:
        return "No active document. Call create_document first."
        
    if paragraph_index is None or paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
        return "Invalid paragraph index."
        
    p = doc.paragraphs[paragraph_index]
    
    try:
        p._element.getparent().remove(p._element)
    except Exception as e:
        return f"Error deleting paragraph: {e}"
        
    if filename:
        doc.save(filename)
    return f"Deleted paragraph {paragraph_index}."

def _hex_to_rgb(hex_str: str):
    try:
        hex_str = hex_str.replace('#', '')
        if len(hex_str) == 6:
            return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))
    except:
        pass
    return None

@mcp.tool()
def create_custom_style(filename: str = None, style_name: str = None, bold: bool = None, italic: bool = None, font_size: int = None, font_name: str = None, color: str = None, base_style: str = 'Normal') -> str:
    """
    Creates a new custom Paragraph style based on an existing style and configures its font properties.
    """
    global current_doc
    doc = current_doc
    if filename and os.path.exists(filename):
        doc = Document(filename)
    if not doc:
        return "No active document. Call create_document first."
        
    if not style_name:
        return "style_name is required."
        
    # Check if style already exists
    try:
        existing = doc.styles[style_name]
        return f"Style '{style_name}' already exists."
    except KeyError:
        pass
        
    # Get base style
    try:
        base = doc.styles[base_style]
    except KeyError:
        return f"Base style '{base_style}' does not exist."
        
    # Create style
    # WD_STYLE_TYPE.PARAGRAPH = 1
    new_style = doc.styles.add_style(style_name, 1)
    new_style.base_style = base
    
    if bold is not None:
        new_style.font.bold = bold
    if italic is not None:
        new_style.font.italic = italic
    if font_size is not None:
        new_style.font.size = Pt(font_size)
    if font_name is not None:
        new_style.font.name = font_name
    if color is not None:
        rgb = _hex_to_rgb(color)
        if rgb:
            new_style.font.color.rgb = rgb
            
    if filename:
        doc.save(filename)
    return f"Created custom paragraph style '{style_name}'."


# Global state for the current document
# In a more complex server, we would use a dictionary mapping session_ids to documents
current_doc = None
current_doc_path = None

def _apply_default_styles(doc):
    """Applies the default styles defined in RULES.md to the document."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.color.rgb = RGBColor(0, 0, 0)
    
    paragraph_format = style.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.line_spacing = 1.15
    paragraph_format.space_after = Pt(12)
    paragraph_format.first_line_indent = Mm(12.7) # 1.27 cm

    # Heading 1 Style
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(16)
    h1.font.bold = False
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.line_spacing = 1.0 # Single spacing for headings per rules
    h1.paragraph_format.first_line_indent = 0

    # Heading 2 Style
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(16) # Defaulting to 16pt as per "Practical Task" rule, user can override
    h2.font.bold = False
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h2.paragraph_format.space_before = Pt(0)
    h2.paragraph_format.space_after = Pt(12)
    h2.paragraph_format.line_spacing = 1.0
    h2.paragraph_format.first_line_indent = 0

@mcp.tool()
def create_document(filename: str = "document.docx") -> str:
    """
    Creates a new compliant Word document with default styles from RULES.md.
    Sets global state to this new document.
    """
    global current_doc, current_doc_path
    current_doc = Document()
    current_doc_path = filename
    
    # Set margins
    sections = current_doc.sections
    for section in sections:
        section.top_margin = Mm(15)
        section.bottom_margin = Mm(15)
        section.left_margin = Mm(20)
        section.right_margin = Mm(20)
        
    _apply_default_styles(current_doc)
    
    return f"Created new document. Ready to save to {filename}."

@mcp.tool()
def add_heading(text: str, level: int = 1) -> str:
    """
    Adds a heading to the document. Defaults to centered, Times New Roman.
    Level 1: 16pt, Normal. Level 2: 16pt, Normal.
    """
    global current_doc
    if not current_doc:
        return "No active document. Call create_document first."
    
    current_doc.add_heading(text, level)
    # If level 2, we might want to allow overriding size, but for now defaults are set in styles.
    # To support specific 14pt Heading 2 (as per rules for "Conclusion to task"), we could check text or add a param.
    return f"Added heading level {level}: '{text}'"

@mcp.tool()
def add_heading_custom(text: str, level: int = 1, font_size: int = None) -> str:
    """
    Adds a heading with optional custom font size (e.g., 14pt for specific Level 2 headings).
    """
    global current_doc
    if not current_doc:
        return "No active document. Call create_document first."
    
    h = current_doc.add_heading(text, level)
    if font_size:
        for run in h.runs:
            run.font.size = Pt(font_size)
    return f"Added custom heading level {level}: '{text}'"

@mcp.tool()
def add_paragraph(text: str, alignment: str = 'JUSTIFY', indent_first_line: bool = True) -> str:
    """
    Adds a paragraph with text.
    Alignment options: LEFT, CENTER, RIGHT, JUSTIFY.
    Default style is Normal (14pt Times New Roman, 1.15 line spacing).
    """
    global current_doc
    if not current_doc:
        return "No active document. Call create_document first."
    
    p = current_doc.add_paragraph(text)
    
    # Force alignment if specified, otherwise inherits from Normal style (JUSTIFY)
    if alignment == 'CENTER':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == 'LEFT':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif alignment == 'RIGHT':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif alignment == 'JUSTIFY':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
    if not indent_first_line:
        p.paragraph_format.first_line_indent = 0
        
    return f"Added paragraph starting with: '{text[:30]}...'"

@mcp.tool()
def add_formatted_text(paragraph_index: int, text: str, bold: bool = False, italic: bool = False, font_size: int = None, lang: str = None) -> str:
    """
    Appends formatted text (a run) to a specific paragraph (by index, -1 for last paragraph).
    Useful for bold keywords or English terms.
    """
    global current_doc
    if not current_doc:
        return "No active document. Call create_document first."
        
    try:
        p = current_doc.paragraphs[paragraph_index]
    except IndexError:
        return f"Paragraph index {paragraph_index} out of range."
        
    run = p.add_run(text)
    run.font.bold = bold
    run.font.italic = italic
    if font_size:
        run.font.size = Pt(font_size)
    if lang:
        # Set language using OXML
        rPr = run._r.get_or_add_rPr()
        lang_element = OxmlElement('w:lang')
        lang_element.set(qn('w:val'), lang)
        rPr.append(lang_element)
        
    return f"Appended formatted text to paragraph {paragraph_index}."

@mcp.tool()
def add_list_item(text: str, style: str = 'List Bullet') -> str:
    """
    Adds a list item using 'List Bullet' or 'List Number'.
    """
    global current_doc
    if not current_doc:
        return "No active document. Call create_document first."
    
    current_doc.add_paragraph(text, style=style)
    return "Added list item."

@mcp.tool()
def add_table(rows: int, cols: int, style: str = 'Table Grid', alignment: str = 'CENTER') -> str:
    """
    Adds a new table to the active document.
    """
    global current_doc
    if not current_doc:
        return "No active document. Call create_document first."
    
    table = current_doc.add_table(rows=rows, cols=cols)
    if style:
        try:
            table.style = style
        except Exception as e:
            return f"Error applying style to table: {e}. Table was created with default style."
    
    # Try applying alignment if supported
    try:
        if alignment == 'CENTER':
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
        elif alignment == 'LEFT':
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
        elif alignment == 'RIGHT':
            table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    except Exception:
        pass # Not all tables support alignment in the same way depending on python-docx version
        
    return f"Added new table ({rows} rows, {cols} columns, style: {style})."

@mcp.tool()
def add_table_row(table_index: int) -> str:
    """
    Appends a new empty row to the specified table.
    """
    global current_doc
    if not current_doc:
        return "No active document. Call create_document first."
    
    try:
        table = current_doc.tables[table_index]
    except IndexError:
        return f"Table index {table_index} out of range."
        
    table.add_row()
    return f"Added new row to table {table_index}. Table now has {len(table.rows)} rows."

@mcp.tool()
def add_table_column(table_index: int, width_pt: int = None) -> str:
    """
    Adds a new column to the specified table.
    """
    global current_doc
    if not current_doc:
        return "No active document. Call create_document first."
    
    try:
        table = current_doc.tables[table_index]
    except IndexError:
        return f"Table index {table_index} out of range."
        
    width = Pt(width_pt) if width_pt else None
    table.add_column(width)
    return f"Added new column to table {table_index}. Table now has {len(table.columns)} columns."

@mcp.tool()
def set_table_cell(table_index: int, row_index: int, col_index: int, text: str, alignment: str = None, bold: bool = False, italic: bool = False) -> str:
    """
    Sets the text of a specific cell and applies basic text formatting and paragraph alignment.
    """
    global current_doc
    if not current_doc:
        return "No active document. Call create_document first."
    
    try:
        table = current_doc.tables[table_index]
        cell = table.cell(row_index, col_index)
    except IndexError:
        return f"Table or cell index out of range."
        
    # Clear existing text and set new text
    cell.text = ""
    # Usually a cell has one empty paragraph when created
    if not cell.paragraphs:
        p = cell.add_paragraph()
    else:
        p = cell.paragraphs[0]
        
    run = p.add_run(text)
    
    if bold:
        run.bold = True
    if italic:
        run.italic = True
        
    if alignment:
        if alignment == 'CENTER':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == 'LEFT':
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alignment == 'RIGHT':
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif alignment == 'JUSTIFY':
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
    return f"Updated cell ({row_index}, {col_index}) in table {table_index}."

@mcp.tool()
def merge_table_cells(table_index: int, start_row: int, start_col: int, end_row: int, end_col: int) -> str:
    """
    Merges a rectangular range of cells into one.
    """
    global current_doc
    if not current_doc:
        return "No active document. Call create_document first."
    
    try:
        table = current_doc.tables[table_index]
        cell_start = table.cell(start_row, start_col)
        cell_end = table.cell(end_row, end_col)
    except IndexError:
        return f"Table or cell indices out of range."
        
    cell_start.merge(cell_end)
    return f"Merged cells from ({start_row}, {start_col}) to ({end_row}, {end_col}) in table {table_index}."

@mcp.tool()
def set_table_cell_style(table_index: int, row_index: int, col_index: int, shading_color: str = None, vertical_alignment: str = None) -> str:
    """
    Applies background color (shading) and vertical alignment to a cell.
    Parameters:
    - shading_color: Hex color string (e.g. 'FF0000') or None.
    - vertical_alignment: 'TOP', 'CENTER', or 'BOTTOM'
    """
    global current_doc
    if not current_doc:
        return "No active document. Call create_document first."
    
    try:
        table = current_doc.tables[table_index]
        cell = table.cell(row_index, col_index)
    except IndexError:
        return f"Table or cell index out of range."

    # Set Vertical Alignment
    if vertical_alignment:
        if vertical_alignment == 'TOP':
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        elif vertical_alignment == 'CENTER':
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        elif vertical_alignment == 'BOTTOM':
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM

    # Set Shading
    if shading_color:
        try:
            # Remove # if present
            shading_color = shading_color.replace('#', '')
            
            tcPr = cell._tc.get_or_add_tcPr()
            
            # Check if shading element already exists
            shdList = tcPr.xpath('w:shd')
            if shdList:
                shdList[0].set(qn('w:fill'), shading_color)
            else:
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), shading_color)
                tcPr.append(shd)
        except Exception as e:
            return f"Error setting shading color: {e}"

    return f"Updated style for cell ({row_index}, {col_index}) in table {table_index}."

@mcp.tool()
def set_table_borders(table_index: int, border_size: int = 4, border_color: str = 'auto') -> str:
    """
    Applies borders to the table via OXML modification.
    border_size is typically 2, 4, 8, etc (in eighths of a point).
    """
    global current_doc
    if not current_doc:
        return "No active document. Call create_document first."
    
    try:
        table = current_doc.tables[table_index]
    except IndexError:
        return f"Table index {table_index} out of range."

    try:
        border_color = border_color.replace('#', '')
        
        tblPr = table._tbl.tblPr
        tblBorders = tblPr.first_child_found_in("w:tblBorders")
        
        if tblBorders is None:
            tblBorders = OxmlElement('w:tblBorders')
            tblPr.append(tblBorders)
            
        # Clean existing borders
        for border in list(tblBorders):
            tblBorders.remove(border)

        val_type = "single"
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), val_type)
            border.set(qn('w:sz'), str(border_size))
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), border_color)
            tblBorders.append(border)
            
    except Exception as e:
        return f"Error setting table borders: {e}"

    return f"Updated borders for table {table_index}."

@mcp.tool()
def delete_table_row(table_index: int, row_index: int) -> str:
    """
    Deletes a specific row from a table.
    """
    global current_doc
    if not current_doc:
        return "No active document. Call create_document first."
    
    try:
        table = current_doc.tables[table_index]
        row = table.rows[row_index]
    except IndexError:
        return f"Table or row index out of range."
        
    try:
        row._element.getparent().remove(row._element)
    except Exception as e:
        return f"Error deleting row: {e}"
        
    return f"Deleted row {row_index} from table {table_index}."

@mcp.tool()
def delete_table_column(table_index: int, col_index: int) -> str:
    """
    Deletes a specific column from a table.
    """
    global current_doc
    if not current_doc:
        return "No active document. Call create_document first."
    
    try:
        table = current_doc.tables[table_index]
        # Test if column index is valid
        if col_index >= len(table.columns):
            return "Column index out of range."
    except IndexError:
        return f"Table index out of range."
        
    try:
        for row in table.rows:
            cell = row.cells[col_index]
            cell._element.getparent().remove(cell._element)
    except Exception as e:
        return f"Error deleting column: {e}"
        
    return f"Deleted column {col_index} from table {table_index}."

@mcp.tool()
def save_document(filename: str = None) -> str:
    """
    Saves the current document to the specified filename (or the one set at creation).
    """
    global current_doc, current_doc_path
    if not current_doc:
        return "No active document. Call create_document first."
        
    target_path = filename if filename else current_doc_path
    if not target_path:
        return "No filename specified."
        
    current_doc.save(target_path)
    return f"Document saved to {target_path}"


# ============================================================================
# DOCUMENT PARAMETERS EXTRACTION FUNCTIONS
# ============================================================================

def _datetime_to_iso(dt_value) -> Optional[str]:
    """Convert datetime object to ISO format string."""
    if dt_value is None:
        return None
    if isinstance(dt_value, datetime):
        return dt_value.isoformat()
    return str(dt_value)


def _rgb_to_hex(rgb_color) -> Optional[str]:
    """Convert RGBColor to hex string."""
    if rgb_color is None:
        return None
    try:
        return str(rgb_color)
    except:
        return None


def _extract_core_properties(doc: Document) -> Dict[str, Any]:
    """
    Extract core properties (metadata) from a document.
    Includes: author, title, subject, keywords, comments, category,
    created, modified, last_modified_by, revision, etc.
    """
    core_props = doc.core_properties
    properties = {}
    
    # Standard core properties - use try/except for compatibility with different python-docx versions
    try:
        properties['author'] = core_props.author
    except AttributeError:
        pass
    
    try:
        properties['title'] = core_props.title
    except AttributeError:
        pass
    
    try:
        properties['subject'] = core_props.subject
    except AttributeError:
        pass
    
    try:
        properties['keywords'] = core_props.keywords
    except AttributeError:
        pass
    
    try:
        properties['comments'] = core_props.comments
    except AttributeError:
        pass
    
    try:
        properties['category'] = core_props.category
    except AttributeError:
        pass
    
    try:
        properties['content_status'] = core_props.content_status
    except AttributeError:
        pass
    
    try:
        properties['identifier'] = core_props.identifier
    except AttributeError:
        pass
    
    try:
        properties['language'] = core_props.language
    except AttributeError:
        pass
    
    try:
        properties['last_modified_by'] = core_props.last_modified_by
    except AttributeError:
        pass
    
    try:
        properties['revision'] = core_props.revision
    except AttributeError:
        pass
    
    # Date properties - convert to ISO format
    try:
        properties['created'] = _datetime_to_iso(core_props.created)
    except AttributeError:
        pass
    
    try:
        properties['modified'] = _datetime_to_iso(core_props.modified)
    except AttributeError:
        pass
    
    try:
        properties['last_printed'] = _datetime_to_iso(core_props.last_printed)
    except AttributeError:
        pass
    
    # Remove None values for cleaner output
    return {k: v for k, v in properties.items() if v is not None}


def _extract_custom_properties(doc: Document) -> Dict[str, Any]:
    """
    Extract custom properties from a document.
    Custom properties are user-defined key-value pairs.
    """
    custom_props = {}
    
    try:
        # Access custom properties through the part
        custom_props_part = doc.part.custom_properties
        if custom_props_part:
            # Parse custom properties XML
            from docx.oxml.ns import qn
            for prop in custom_props_part.element:
                prop_name = prop.get('name')
                if prop_name:
                    # Get the value element
                    value_elem = prop.find(qn('vt:lpwstr'))
                    if value_elem is not None:
                        custom_props[prop_name] = value_elem.text
                    else:
                        # Try other value types
                        for child in prop:
                            if child.text:
                                custom_props[prop_name] = child.text
                                break
    except Exception as e:
        # Custom properties might not exist
        pass
    
    return custom_props


def _extract_document_variables(doc: Document) -> Dict[str, str]:
    """
    Extract document variables from a document.
    Document variables are used for mail merge and document automation.
    """
    variables = {}
    
    try:
        # Document variables are stored in settings.xml
        settings = doc.settings.element
        
        # Find all docVar elements
        for var in settings.iter(qn('w:docVar')):
            name = var.get(qn('w:name'))
            value = var.get(qn('w:val'))
            if name:
                variables[name] = value if value else ''
    except Exception as e:
        pass
    
    return variables


def _extract_section_properties(doc: Document) -> List[Dict[str, Any]]:
    """
    Extract section properties including margins, page size, orientation.
    """
    sections = []
    
    for i, section in enumerate(doc.sections):
        section_info = {
            'index': i,
            'page_width': None,
            'page_height': None,
            'orientation': 'portrait',
            'margins': {},
            'header_footer': {}
        }
        
        # Page dimensions
        try:
            section_info['page_width'] = section.page_width.pt if section.page_width else None
            section_info['page_height'] = section.page_height.pt if section.page_height else None
        except:
            pass
        
        # Orientation
        try:
            if section.orientation == WD_ORIENT.LANDSCAPE:
                section_info['orientation'] = 'landscape'
        except:
            pass
        
        # Margins
        try:
            margins = {}
            if section.top_margin:
                margins['top_mm'] = section.top_margin.mm
                margins['top_pt'] = section.top_margin.pt
            if section.bottom_margin:
                margins['bottom_mm'] = section.bottom_margin.mm
                margins['bottom_pt'] = section.bottom_margin.pt
            if section.left_margin:
                margins['left_mm'] = section.left_margin.mm
                margins['left_pt'] = section.left_margin.pt
            if section.right_margin:
                margins['right_mm'] = section.right_margin.mm
                margins['right_pt'] = section.right_margin.pt
            if section.gutter:
                margins['gutter_mm'] = section.gutter.mm
                margins['gutter_pt'] = section.gutter.pt
            section_info['margins'] = margins
        except:
            pass
        
        # Header/Footer distances
        try:
            header_footer = {}
            if section.header_distance:
                header_footer['header_distance_mm'] = section.header_distance.mm
            if section.footer_distance:
                header_footer['footer_distance_mm'] = section.footer_distance.mm
            section_info['header_footer'] = header_footer
        except:
            pass
        
        # Different first page
        try:
            section_info['different_first_page'] = section.different_first_page_header_footer
        except:
            pass
        
        sections.append(section_info)
    
    return sections


def _extract_styles_info(doc: Document, only_used: bool = True, exclude_hidden: bool = True) -> Dict[str, Any]:
    """
    Extract style information from a document.
    Returns information about paragraph and character styles.
    
    Args:
        doc: The Document object
        only_used: If True, only extract styles that are actually used in the document
        exclude_hidden: If True, exclude hidden styles
    """
    styles_info = {
        'paragraph_styles': {},
        'character_styles': {},
        'table_styles': {},
        'list_styles': {}
    }
    
    # Find styles that are actually used in the document
    used_styles = set()
    if only_used:
        for para in doc.paragraphs:
            if para.style:
                used_styles.add(para.style.name)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.style:
                            used_styles.add(para.style.name)
    
    for style in doc.styles:
        # Skip hidden styles if requested
        if exclude_hidden:
            try:
                if style.hidden:
                    continue
            except:
                pass
        
        # Skip unused styles if only_used is True
        if only_used and style.name not in used_styles:
            continue
        
        style_info = {}
        
        # Basic style properties - only store non-default values
        try:
            if style.style_id and style.style_id != style.name:
                style_info['style_id'] = style.style_id
        except:
            pass
        
        # Font properties - only store non-default/important values
        try:
            font = style.font
            font_info = {}
            if font.name:
                font_info['name'] = font.name
            if font.size:
                font_info['size_pt'] = font.size.pt
            if font.bold is True:  # Only store if True, not False or None
                font_info['bold'] = font.bold
            if font.italic is True:  # Only store if True
                font_info['italic'] = font.italic
            if font.underline is not None and font.underline != False:
                font_info['underline'] = str(font.underline)
            if font.color.rgb:
                font_info['color'] = _rgb_to_hex(font.color.rgb)
            if font_info:  # Only add if not empty
                style_info['font'] = font_info
        except:
            pass
        
        # Paragraph format properties - only store non-default values
        try:
            pf = style.paragraph_format
            para_info = {}
            if pf.alignment:
                para_info['alignment'] = str(pf.alignment)
            if pf.line_spacing and pf.line_spacing != 1.0:
                para_info['line_spacing'] = pf.line_spacing
            if pf.space_before and pf.space_before.pt != 0:
                para_info['space_before_pt'] = pf.space_before.pt
            if pf.space_after and pf.space_after.pt != 0:
                para_info['space_after_pt'] = pf.space_after.pt
            if pf.first_line_indent and pf.first_line_indent.mm != 0:
                para_info['first_line_indent_mm'] = pf.first_line_indent.mm
            if pf.left_indent and pf.left_indent.mm != 0:
                para_info['left_indent_mm'] = pf.left_indent.mm
            if pf.right_indent and pf.right_indent.mm != 0:
                para_info['right_indent_mm'] = pf.right_indent.mm
            if para_info:  # Only add if not empty
                style_info['paragraph_format'] = para_info
        except:
            pass
        
        # Categorize by style type
        if style.type == WD_STYLE_TYPE.PARAGRAPH:
            styles_info['paragraph_styles'][style.name] = style_info
        elif style.type == WD_STYLE_TYPE.CHARACTER:
            styles_info['character_styles'][style.name] = style_info
        elif style.type == WD_STYLE_TYPE.TABLE:
            styles_info['table_styles'][style.name] = style_info
        elif style.type == WD_STYLE_TYPE.LIST:
            styles_info['list_styles'][style.name] = style_info
    
    return styles_info


def _extract_numbering_info(doc: Document) -> Dict[str, Any]:
    """
    Extract numbering/bullet list definitions from a document.
    """
    numbering_info = {}
    
    try:
        # Access numbering part
        numbering_part = doc.part.numbering_part
        if numbering_part:
            # Parse numbering definitions
            for num in numbering_part.element.iter(qn('w:num')):
                num_id = num.get(qn('w:numId'))
                abstract_num_id = None
                for ref in num.iter(qn('w:abstractNumId')):
                    abstract_num_id = ref.get(qn('w:val'))
                    break
                if num_id:
                    numbering_info[num_id] = {'abstract_num_id': abstract_num_id}
    except:
        pass
    
    return numbering_info


def _extract_headers_footers(doc: Document) -> Dict[str, Any]:
    """
    Extract headers and footers content from a document.
    """
    headers_footers = {'headers': {}, 'footers': {}}
    
    for i, section in enumerate(doc.sections):
        section_key = f'section_{i}'
        
        # Headers
        try:
            header = section.header
            if header and not header.is_linked_to_previous:
                text = '\n'.join([p.text for p in header.paragraphs if p.text.strip()])
                if text:
                    headers_footers['headers'][section_key] = text
        except:
            pass
        
        try:
            first_header = section.first_page_header
            if first_header and not first_header.is_linked_to_previous:
                text = '\n'.join([p.text for p in first_header.paragraphs if p.text.strip()])
                if text:
                    headers_footers['headers'][f'{section_key}_first'] = text
        except:
            pass
        
        # Footers
        try:
            footer = section.footer
            if footer and not footer.is_linked_to_previous:
                text = '\n'.join([p.text for p in footer.paragraphs if p.text.strip()])
                if text:
                    headers_footers['footers'][section_key] = text
        except:
            pass
        
        try:
            first_footer = section.first_page_footer
            if first_footer and not first_footer.is_linked_to_previous:
                text = '\n'.join([p.text for p in first_footer.paragraphs if p.text.strip()])
                if text:
                    headers_footers['footers'][f'{section_key}_first'] = text
        except:
            pass
    
    return headers_footers


def _extract_tables_info(doc: Document) -> List[Dict[str, Any]]:
    """
    Extract table structure information from a document.
    """
    tables = []
    
    for i, table in enumerate(doc.tables):
        table_info = {
            'index': i,
            'rows': len(table.rows),
            'columns': len(table.columns),
            'cells': []
        }
        
        # Extract cell information
        for row_idx, row in enumerate(table.rows):
            row_cells = []
            for col_idx, cell in enumerate(row.cells):
                cell_info = {
                    'row': row_idx,
                    'col': col_idx,
                    'text': cell.text[:100] + '...' if len(cell.text) > 100 else cell.text
                }
                row_cells.append(cell_info)
            table_info['cells'].append(row_cells)
        
        tables.append(table_info)
    
    return tables


@mcp.tool()
def extract_document_parameters(filename: str = None, compact: bool = False, all_styles: bool = False) -> str:
    """
    Extract all document parameters from the current or specified document.
    Returns a JSON string with core properties, custom properties, document variables,
    section properties, styles, and other metadata.
    
    Args:
        filename: Optional path to a DOCX file. If not provided, uses the current document.
        compact: If True, output compact JSON without indentation (smaller size).
        all_styles: If True, extract all styles including unused ones.
    
    Returns:
        JSON string containing all extracted document parameters.
    """
    global current_doc, current_doc_path
    
    doc = None
    if filename:
        if not os.path.exists(filename):
            return json.dumps({'error': f'File not found: {filename}'})
        doc = Document(filename)
    elif current_doc:
        doc = current_doc
    else:
        return json.dumps({'error': 'No document loaded. Provide a filename or create a document first.'})
    
    parameters = {
        'source_file': filename if filename else current_doc_path,
        'extraction_timestamp': datetime.now().isoformat(),
        'core_properties': _extract_core_properties(doc),
        'custom_properties': _extract_custom_properties(doc),
        'document_variables': _extract_document_variables(doc),
        'sections': _extract_section_properties(doc),
        'styles': _extract_styles_info(doc, only_used=not all_styles),
        'numbering': _extract_numbering_info(doc),
        'headers_footers': _extract_headers_footers(doc),
        'tables': _extract_tables_info(doc),
        'paragraphs_count': len(doc.paragraphs),
        'tables_count': len(doc.tables)
    }
    
    if compact:
        return json.dumps(parameters, ensure_ascii=False)
    else:
        return json.dumps(parameters, indent=2, ensure_ascii=False)


@mcp.tool()
def extract_core_properties(filename: str = None) -> str:
    """
    Extract only core properties (metadata) from a document.
    
    Args:
        filename: Optional path to a DOCX file. If not provided, uses the current document.
    
    Returns:
        JSON string containing core properties.
    """
    global current_doc, current_doc_path
    
    doc = None
    if filename:
        if not os.path.exists(filename):
            return json.dumps({'error': f'File not found: {filename}'})
        doc = Document(filename)
    elif current_doc:
        doc = current_doc
    else:
        return json.dumps({'error': 'No document loaded.'})
    
    return json.dumps(_extract_core_properties(doc), indent=2, ensure_ascii=False)


@mcp.tool()
def extract_custom_properties(filename: str = None) -> str:
    """
    Extract only custom properties from a document.
    
    Args:
        filename: Optional path to a DOCX file. If not provided, uses the current document.
    
    Returns:
        JSON string containing custom properties.
    """
    global current_doc, current_doc_path
    
    doc = None
    if filename:
        if not os.path.exists(filename):
            return json.dumps({'error': f'File not found: {filename}'})
        doc = Document(filename)
    elif current_doc:
        doc = current_doc
    else:
        return json.dumps({'error': 'No document loaded.'})
    
    return json.dumps(_extract_custom_properties(doc), indent=2, ensure_ascii=False)


@mcp.tool()
def extract_document_variables(filename: str = None) -> str:
    """
    Extract only document variables from a document.
    
    Args:
        filename: Optional path to a DOCX file. If not provided, uses the current document.
    
    Returns:
        JSON string containing document variables.
    """
    global current_doc, current_doc_path
    
    doc = None
    if filename:
        if not os.path.exists(filename):
            return json.dumps({'error': f'File not found: {filename}'})
        doc = Document(filename)
    elif current_doc:
        doc = current_doc
    else:
        return json.dumps({'error': 'No document loaded.'})
    
    return json.dumps(_extract_document_variables(doc), indent=2, ensure_ascii=False)


@mcp.tool()
def extract_section_properties(filename: str = None) -> str:
    """
    Extract section properties (margins, page size, orientation) from a document.
    
    Args:
        filename: Optional path to a DOCX file. If not provided, uses the current document.
    
    Returns:
        JSON string containing section properties.
    """
    global current_doc, current_doc_path
    
    doc = None
    if filename:
        if not os.path.exists(filename):
            return json.dumps({'error': f'File not found: {filename}'})
        doc = Document(filename)
    elif current_doc:
        doc = current_doc
    else:
        return json.dumps({'error': 'No document loaded.'})
    
    return json.dumps(_extract_section_properties(doc), indent=2, ensure_ascii=False)


@mcp.tool()
def extract_styles_info(filename: str = None, all_styles: bool = False, compact: bool = False) -> str:
    """
    Extract style information from a document.
    
    Args:
        filename: Optional path to a DOCX file. If not provided, uses the current document.
        all_styles: If True, extract all styles including unused ones.
        compact: If True, output compact JSON without indentation.
    
    Returns:
        JSON string containing style information.
    """
    global current_doc, current_doc_path
    
    doc = None
    if filename:
        if not os.path.exists(filename):
            return json.dumps({'error': f'File not found: {filename}'})
        doc = Document(filename)
    elif current_doc:
        doc = current_doc
    else:
        return json.dumps({'error': 'No document loaded.'})
    
    styles = _extract_styles_info(doc, only_used=not all_styles)
    
    if compact:
        return json.dumps(styles, ensure_ascii=False)
    else:
        return json.dumps(styles, indent=2, ensure_ascii=False)


@mcp.tool()
def load_template(filename: str) -> str:
    """
    Load an existing document as a template for further modifications.
    Sets the loaded document as the current document.
    
    Args:
        filename: Path to the DOCX template file.
    
    Returns:
        Status message.
    """
    global current_doc, current_doc_path
    
    if not os.path.exists(filename):
        return f"Error: File not found: {filename}"
    
    current_doc = Document(filename)
    current_doc_path = filename
    
    return f"Template loaded from {filename}. Document is ready for modifications."


def _apply_styles_from_params(doc: Document, styles_params: Dict[str, Any]) -> int:
    """
    Apply styles from parameters to a document.
    
    Args:
        doc: The Document object to apply styles to
        styles_params: Dictionary containing style parameters
    
    Returns:
        Number of styles applied
    """
    applied_count = 0
    
    # Helper function to parse alignment string
    def parse_alignment(align_str: str):
        if not align_str:
            return None
        align_str = align_str.upper()
        if 'CENTER' in align_str:
            return WD_ALIGN_PARAGRAPH.CENTER
        elif 'LEFT' in align_str:
            return WD_ALIGN_PARAGRAPH.LEFT
        elif 'RIGHT' in align_str:
            return WD_ALIGN_PARAGRAPH.RIGHT
        elif 'JUSTIFY' in align_str:
            return WD_ALIGN_PARAGRAPH.JUSTIFY
        return None
    
    # Helper function to parse color hex to RGBColor
    def parse_color(color_str: str):
        if not color_str:
            return None
        try:
            # Remove # if present
            color_str = color_str.replace('#', '')
            if len(color_str) == 6:
                return RGBColor(int(color_str[0:2], 16), int(color_str[2:4], 16), int(color_str[4:6], 16))
        except:
            pass
        return None
    
    # Process paragraph styles
    if 'paragraph_styles' in styles_params:
        for style_name, style_info in styles_params['paragraph_styles'].items():
            try:
                # Try to get existing style
                try:
                    style = doc.styles[style_name]
                except KeyError:
                    # Style doesn't exist, skip
                    continue
                
                # Apply font properties
                if 'font' in style_info:
                    font_info = style_info['font']
                    font = style.font
                    
                    if 'name' in font_info:
                        font.name = font_info['name']
                    if 'size_pt' in font_info:
                        font.size = Pt(font_info['size_pt'])
                    if 'bold' in font_info:
                        font.bold = font_info['bold']
                    if 'italic' in font_info:
                        font.italic = font_info['italic']
                    if 'underline' in font_info:
                        # Handle underline - could be string like "True" or actual value
                        underline_val = font_info['underline']
                        if isinstance(underline_val, bool):
                            font.underline = underline_val
                        elif isinstance(underline_val, str):
                            font.underline = underline_val.lower() == 'true'
                    if 'color' in font_info:
                        color = parse_color(font_info['color'])
                        if color:
                            font.color.rgb = color
                
                # Apply paragraph format properties
                if 'paragraph_format' in style_info:
                    para_info = style_info['paragraph_format']
                    pf = style.paragraph_format
                    
                    if 'alignment' in para_info:
                        alignment = parse_alignment(para_info['alignment'])
                        if alignment:
                            pf.alignment = alignment
                    
                    if 'line_spacing' in para_info:
                        pf.line_spacing = para_info['line_spacing']
                    
                    if 'space_before_pt' in para_info:
                        pf.space_before = Pt(para_info['space_before_pt'])
                    
                    if 'space_after_pt' in para_info:
                        pf.space_after = Pt(para_info['space_after_pt'])
                    
                    if 'first_line_indent_mm' in para_info:
                        pf.first_line_indent = Mm(para_info['first_line_indent_mm'])
                    
                    if 'left_indent_mm' in para_info:
                        pf.left_indent = Mm(para_info['left_indent_mm'])
                    
                    if 'right_indent_mm' in para_info:
                        pf.right_indent = Mm(para_info['right_indent_mm'])
                
                applied_count += 1
                
            except Exception as e:
                # Skip styles that can't be applied
                continue
    
    return applied_count


@mcp.tool()
def apply_template_parameters(parameters_json: str, output_filename: str = None) -> str:
    """
    Apply extracted parameters to create a new document based on template settings.
    This function takes a JSON string of parameters (from extract_document_parameters)
    and creates a new document with those settings applied.
    
    Args:
        parameters_json: JSON string containing document parameters.
        output_filename: Optional filename for the new document.
    
    Returns:
        Status message.
    """
    global current_doc, current_doc_path
    
    try:
        params = json.loads(parameters_json)
    except json.JSONDecodeError as e:
        return f"Error parsing JSON: {e}"
    
    # Create new document
    current_doc = Document()
    current_doc_path = output_filename if output_filename else "new_document.docx"
    
    styles_applied = 0
    
    # Apply section properties
    if 'sections' in params:
        for i, section_params in enumerate(params['sections']):
            if i < len(current_doc.sections):
                section = current_doc.sections[i]
                
                # Apply margins
                if 'margins' in section_params:
                    margins = section_params['margins']
                    if 'top_mm' in margins:
                        section.top_margin = Mm(margins['top_mm'])
                    if 'bottom_mm' in margins:
                        section.bottom_margin = Mm(margins['bottom_mm'])
                    if 'left_mm' in margins:
                        section.left_margin = Mm(margins['left_mm'])
                    if 'right_mm' in margins:
                        section.right_margin = Mm(margins['right_mm'])
                
                # Apply page size
                if 'page_width' in section_params and 'page_height' in section_params:
                    section.page_width = Pt(section_params['page_width'])
                    section.page_height = Pt(section_params['page_height'])
                
                # Apply orientation
                if 'orientation' in section_params:
                    if section_params['orientation'] == 'landscape':
                        section.orientation = WD_ORIENT.LANDSCAPE
    
    # Apply styles
    if 'styles' in params:
        styles_applied = _apply_styles_from_params(current_doc, params['styles'])
    
    # Apply core properties
    if 'core_properties' in params:
        core_props = params['core_properties']
        cp = current_doc.core_properties
        
        if 'author' in core_props:
            cp.author = core_props['author']
        if 'title' in core_props:
            cp.title = core_props['title']
        if 'subject' in core_props:
            cp.subject = core_props['subject']
        if 'keywords' in core_props:
            cp.keywords = core_props['keywords']
        if 'comments' in core_props:
            cp.comments = core_props['comments']
        if 'category' in core_props:
            cp.category = core_props['category']
    
    return f"Created new document with template parameters. Styles applied: {styles_applied}. Ready to save to {current_doc_path}."


@mcp.tool()
def set_core_property(property_name: str, value: str) -> str:
    """
    Set a core property on the current document.
    
    Args:
        property_name: Name of the property (author, title, subject, keywords, comments, category).
        value: Value to set for the property.
    
    Returns:
        Status message.
    """
    global current_doc
    
    if not current_doc:
        return "No active document. Call create_document or load_template first."
    
    valid_properties = ['author', 'title', 'subject', 'keywords', 'comments', 'category', 
                        'content_status', 'identifier', 'language', 'status', 'version']
    
    if property_name not in valid_properties:
        return f"Invalid property. Valid properties are: {', '.join(valid_properties)}"
    
    setattr(current_doc.core_properties, property_name, value)
    return f"Set {property_name} to '{value}'"


@mcp.tool()
def set_custom_property(property_name: str, value: str) -> str:
    """
    Set a custom property on the current document.
    
    Args:
        property_name: Name of the custom property.
        value: Value to set for the property.
    
    Returns:
        Status message.
    """
    global current_doc
    
    if not current_doc:
        return "No active document. Call create_document or load_template first."
    
    try:
        # python-docx has limited support for custom properties
        # We need to work with the underlying XML
        from docx.oxml.ns import qn
        
        # Get or create custom properties part
        custom_props = current_doc.part.custom_properties
        if custom_props is None:
            # Create custom properties part
            from docx.parts.custom_properties import CustomPropertiesPart
            custom_props = CustomPropertiesPart()
            current_doc.part._custom_properties = custom_props
        
        # This is a simplified implementation
        # Full implementation would require more XML manipulation
        return f"Custom property '{property_name}' noted. Note: Full custom property support requires additional XML handling."
    except Exception as e:
        return f"Error setting custom property: {e}"


@mcp.tool()
def get_document_structure(filename: str = None) -> str:
    """
    Get a summary of the document structure including headings, paragraphs, and tables.
    
    Args:
        filename: Optional path to a DOCX file. If not provided, uses the current document.
    
    Returns:
        JSON string containing document structure.
    """
    global current_doc, current_doc_path
    
    doc = None
    if filename:
        if not os.path.exists(filename):
            return json.dumps({'error': f'File not found: {filename}'})
        doc = Document(filename)
    elif current_doc:
        doc = current_doc
    else:
        return json.dumps({'error': 'No document loaded.'})
    
    structure = {
        'headings': [],
        'paragraphs': [],
        'tables_count': len(doc.tables)
    }
    
    for i, para in enumerate(doc.paragraphs):
        if para.style.name.startswith('Heading'):
            structure['headings'].append({
                'index': i,
                'level': para.style.name,
                'text': para.text[:100] + '...' if len(para.text) > 100 else para.text
            })
        elif para.text.strip():
            structure['paragraphs'].append({
                'index': i,
                'style': para.style.name,
                'text_preview': para.text[:50] + '...' if len(para.text) > 50 else para.text
            })
    
    return json.dumps(structure, indent=2, ensure_ascii=False)


if __name__ == "__main__":
    mcp.run()
